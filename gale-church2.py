import json
import numpy as np
from nltk.translate.gale_church import align_blocks
from docx import Document

# Load sentences.json
with open('sentences.json', 'r', encoding='utf-8') as f:
    sentence_pairs = json.load(f)

# Extract Sinhala and English sentences
sinhala_sentences = [pair["Sinhala_Sentence"] for pair in sentence_pairs]
english_sentences = [pair["English_Sentence"] for pair in sentence_pairs]


# Calculate sentence lengths
sinhala_lengths = np.array([len(sentence) for sentence in sinhala_sentences])
english_lengths = np.array([len(sentence) for sentence in english_sentences])

# Align sentences using Gale-Church algorithm
def gale_church_realign(sinhala_lengths, english_lengths):
    # Align sentences based on their lengths
    alignments = align_blocks(sinhala_lengths.tolist(), english_lengths.tolist())

    
    # Create a dictionary to map sentence indices to alignments
    aligned_sentences = []
    used_sinhala_indices = set()
    used_english_indices = set()
    
    for (sin_idx, eng_idx) in alignments:
        if sin_idx not in used_sinhala_indices and eng_idx not in used_english_indices:
            aligned_sentences.append({
                'Sinhala_Sentence': sinhala_sentences[sin_idx],
                'English_Sentence': english_sentences[eng_idx]
            })
            used_sinhala_indices.add(sin_idx)
            used_english_indices.add(eng_idx)
    
    return aligned_sentences

# Perform alignment
aligned_sentences = gale_church_realign(sinhala_lengths, english_lengths)

# Save the aligned sentences to a new JSON file
with open('re_aligned_sentences.json', 'w', encoding='utf-8') as f:
    json.dump(aligned_sentences, f, ensure_ascii=False, indent=4)

# Save the aligned sentences to a Word document
def save_to_word(aligned_sentences):
    doc = Document()
    doc.add_heading('Re-Aligned Sinhala-English Sentences', 0)

    for pair in aligned_sentences:
        doc.add_paragraph(f"Sinhala Sentence: {pair['Sinhala_Sentence']}")
        doc.add_paragraph(f"English Sentence: {pair['English_Sentence']}")
        doc.add_paragraph('------------------------------')

    doc.save('re_aligned_sentences.docx')

# Save to Word document
save_to_word(aligned_sentences)
